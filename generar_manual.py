#!/usr/bin/env python3
"""
Generador del Manual de Usuario — NSG Cotizador
Genera un archivo .docx profesional con python-docx.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda de tabla."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Agrega tabla con encabezado gris y bordes."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Headers
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '005EC2')
    # Rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

def add_bullet_list(doc, items, level=0):
    """Agrega lista viñeteada."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)

def add_numbered_list(doc, items):
    """Agrega lista numerada."""
    for item in items:
        doc.add_paragraph(item, style='List Number')

def add_note(doc, text, prefix="Nota"):
    """Agrega párrafo de nota con estilo diferenciado."""
    p = doc.add_paragraph()
    run = p.add_run(f"{prefix}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x5E, 0xC2)
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.italic = True

def add_tip(doc, text):
    """Agrega consejo."""
    p = doc.add_paragraph()
    run = p.add_run("Consejo: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x30, 0xD1, 0x58)
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.italic = True

def add_warning(doc, text):
    """Agrega advertencia."""
    p = doc.add_paragraph()
    run = p.add_run("Advertencia: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x45, 0x3A)
    run.font.size = Pt(9)
    run2 = p.add_run(text)
    run2.font.size = Pt(9)
    run2.font.italic = True


# ─── Documento principal ────────────────────────────────────────────────────

def generar_manual():
    doc = Document()

    # ── Estilos personalizados ──────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)
        hs.font.name = 'Calibri'

    # ── PORTADA ─────────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NSG COTIZADOR")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x5E, 0xC2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual de Usuario")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1D, 0x1D, 0x1F)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Sistema de Cotización Automatizada de Equipos de Cómputo")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NSG Latinoamérica E.I.R.L.")
    run.font.size = Pt(12)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versión 1.0 — Mayo 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6E, 0x6E, 0x73)

    doc.add_page_break()

    # ── TABLA DE CONTENIDO ──────────────────────────────────────────────────
    doc.add_heading('Tabla de Contenido', level=1)
    p = doc.add_paragraph('Este manual contiene las instrucciones completas para utilizar el sistema NSG Cotizador. Utilice la tabla de contenido de su lector de documentos para navegar entre secciones.')
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCCIÓN
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('1. Introducción', level=1)

    doc.add_heading('1.1 Propósito del sistema', level=2)
    doc.add_paragraph(
        'NSG Cotizador es un sistema web diseñado para agilizar el proceso de cotización de equipos de cómputo personalizados. '
        'Permite armar configuraciones de PC paso a paso seleccionando componentes compatibles, calcular precios en USD y PEN '
        'con márgenes e IGV configurables, generar cotizaciones en PDF y Excel, y administrar el catálogo de productos.'
    )

    doc.add_heading('1.2 Audiencia', level=2)
    add_bullet_list(doc, [
        'Administradores: gestionan catálogo, configuración, importación CSV y validación de tickets.',
        'Usuarios registrados: arman configuraciones, generan cotizaciones y consultan su historial.',
        'Invitados: pueden explorar el cotizador y armar configuraciones, pero necesitan registrarse para generar cotizaciones.'
    ])

    doc.add_heading('1.3 Características principales', level=2)
    add_bullet_list(doc, [
        'Cotizador paso a paso con 7 categorías de componentes principales y 8 subcategorías de extras.',
        'Validación automática de compatibilidad entre componentes (socket, tipo RAM, form factor, etc.).',
        'Generación de cotizaciones en PDF comercial, PDF técnico y Excel.',
        'Asistente de IA integrado para recomendaciones de configuración.',
        'Importación masiva de productos desde archivos CSV con enriquecimiento automático por IA.',
        'Dashboard de métricas operativas en tiempo real.',
        'Tipo de cambio USD/PEN automático o manual.',
        'Sistema de favoritos para guardar productos de referencia.',
        'Compartir configuraciones mediante enlaces.',
        'Notificaciones por correo electrónico al cliente.',
        'Soporte completo para modo oscuro y accesibilidad WCAG AA.'
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. ACCESO Y REGISTRO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('2. Acceso y Registro', level=1)

    doc.add_heading('2.1 Crear cuenta nueva', level=2)
    doc.add_paragraph(
        'Para crear una cuenta, acceda a la página de registro desde el enlace "Regístrate" en la página de inicio de sesión.'
    )
    doc.add_heading('Campos del formulario', level=3)
    add_styled_table(doc,
        ['Campo', 'Requerido', 'Validación'],
        [
            ['Usuario', 'Sí', 'Mínimo 3 caracteres. Solo letras, números y guion bajo.'],
            ['Nombre completo', 'Sí', 'Mínimo 2 caracteres.'],
            ['Correo electrónico', 'Sí', 'Formato válido (ejemplo@dominio.com).'],
            ['Teléfono', 'No', 'Opcional.'],
            ['Contraseña', 'Sí', 'Mínimo 8 caracteres, al menos: 1 mayúscula, 1 minúscula, 1 número, 1 carácter especial.'],
            ['Confirmar contraseña', 'Sí', 'Debe coincidir exactamente con la contraseña.'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'El sistema incluye un medidor visual de fortaleza de contraseña que cambia de color progresivamente '
        '(rojo → amarillo → verde) mientras escribe. Al completar el registro exitoso, será redirigido al Cotizador.'
    )

    doc.add_heading('2.2 Iniciar sesión', level=2)
    doc.add_paragraph(
        'Ingrese su usuario y contraseña en la página de inicio de sesión. Puede alternar la visibilidad '
        'de la contraseña con el ícono de ojo junto al campo. Si las credenciales son correctas, el sistema '
        'lo redirigirá al panel de administración de productos (si es admin) o al cotizador (si es usuario).'
    )
    add_tip(doc, 'Use el enlace "Recuperar contraseña" si olvidó sus credenciales.')

    doc.add_heading('2.3 Recuperar contraseña', level=2)
    doc.add_paragraph(
        'Si olvidó su contraseña, puede solicitar un enlace de recuperación por dos métodos:'
    )
    add_numbered_list(doc, [
        'Correo electrónico: ingrese su correo registrado y recibirá un enlace de restablecimiento. '
        'El enlace expira en 5 minutos.',
        'Número de teléfono: ingrese su número registrado (7-15 dígitos) y el sistema enviará '
        'el enlace al correo asociado a ese teléfono.',
    ])
    add_note(doc, 'Revise también la carpeta de correo no deseado (spam) si no recibe el correo.')

    doc.add_heading('2.4 Restablecer contraseña', level=2)
    doc.add_paragraph(
        'Al hacer clic en el enlace de recuperación recibido por correo, accederá a la página de restablecimiento. '
        'Ingrese su nueva contraseña (sujeta a las mismas reglas de fortaleza del registro) y confírmela. '
        'Si el enlace es inválido o ha expirado, el sistema mostrará un mensaje de error y le ofrecerá solicitar uno nuevo.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. COTIZADOR DE PC
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('3. Cotizador de PC', level=1)

    doc.add_heading('3.1 Vista general', level=2)
    doc.add_paragraph(
        'El Cotizador es la página principal del sistema. Permite armar una configuración de PC paso a paso, '
        'seleccionando componentes de 7 categorías obligatorias y agregando extras opcionales. '
        'El sistema valida la compatibilidad entre componentes automáticamente y calcula el precio total '
        'con margen, IGV y tipo de cambio aplicados.'
    )

    doc.add_heading('3.2 Pasos del cotizador', level=2)
    doc.add_paragraph(
        'El proceso sigue una secuencia obligatoria. Cada paso debe completarse antes de avanzar al siguiente:'
    )
    add_styled_table(doc,
        ['Paso', 'Categoría', 'Obligatorio', 'Notas'],
        [
            ['1', 'Procesador', 'Sí', 'Se filtran placas madre compatibles por socket.'],
            ['2', 'Placa Madre', 'Sí', 'Se filtran RAM y almacenamiento compatibles.'],
            ['3', 'Memoria RAM', 'Sí', 'Puede agregar múltiples módulos.'],
            ['4', 'Almacenamiento', 'Sí', 'SSD, HDD, M.2/NVMe filtrados por placa.'],
            ['5', 'Tarjeta Gráfica', 'Sí*', 'Opcional si el procesador tiene gráficos integrados.'],
            ['6', 'Fuente de Poder', 'Sí', 'Filtrada por requisitos de GPU.'],
            ['7', 'Case/Gabinete', 'Sí', 'Filtrado por form factor de placa madre.'],
            ['8', 'Otros/Extras', 'No', 'Periféricos, audio, software, monitor, etc.'],
        ]
    )

    doc.add_heading('3.3 Filtros de productos', level=2)
    doc.add_paragraph('Cada paso ofrece filtros específicos según la categoría:')
    add_styled_table(doc,
        ['Categoría', 'Filtros disponibles'],
        [
            ['Procesador', 'Marca, Modelo (Core i3/i5/i7/i9, Ryzen 3/5/7/9)'],
            ['Placa Madre', 'Marca'],
            ['RAM', 'Marca, Frecuencia (MHz), Capacidad (GB)'],
            ['Almacenamiento', 'Marca, Capacidad (GB/TB), Tipo (M.2/NVMe, SSD, HDD)'],
            ['GPU', 'Marca, Serie (RTX 30/40, RX 6000/7000, GTX)'],
            ['Fuente', 'Marca, Potencia (W)'],
            ['Case', 'Marca'],
        ]
    )

    doc.add_heading('3.4 Vista compacta vs detallada', level=2)
    doc.add_paragraph(
        'Use el selector "Vista compacta / Vista detallada" en la parte superior para alternar entre:'
    )
    add_bullet_list(doc, [
        'Vista compacta: lista rápida con nombre, precio, stock y acciones básicas.',
        'Vista detallada: tarjetas con especificaciones técnicas completas, descripción y badges de estado.'
    ])

    doc.add_heading('3.5 Filtro de disponibilidad', level=2)
    doc.add_paragraph(
        'El interruptor "Solo disponibles" filtra los productos para mostrar únicamente aquellos con stock mayor a 0 '
        'o marcados como "disponible a pedido". Al desactivarlo, se muestran todos los productos del catálogo.'
    )
    doc.add_paragraph('Indicadores de stock en cada producto:')
    add_bullet_list(doc, [
        'Stock N (verde): unidades disponibles en almacén.',
        'A pedido (amarillo): disponible bajo pedido con tiempo de entrega indicado en días.',
        'Sin stock (rojo): producto agotado temporalmente.'
    ])

    doc.add_heading('3.6 Ordenar por precio', level=2)
    doc.add_paragraph(
        'Use el selector de ordenamiento para listar productos de tres formas:'
    )
    add_bullet_list(doc, [
        'Relevancia (predeterminado): orden original del catálogo.',
        'Menor precio: del más económico al más costoso.',
        'Mayor precio: del más costoso al más económico.'
    ])

    doc.add_heading('3.7 Validación automática de compatibilidad', level=2)
    doc.add_paragraph(
        'Al seleccionar un componente, el sistema valida automáticamente la compatibilidad con los componentes '
        'ya seleccionados. Las reglas de compatibilidad incluyen:'
    )
    add_bullet_list(doc, [
        'Socket del procesador debe coincidir con el de la placa madre.',
        'Tipo de RAM (DDR4/DDR5) debe coincidir con el soportado por la placa madre.',
        'Almacenamiento M.2/NVMe solo se muestra si la placa madre lo soporta.',
        'Form factor de la placa madre debe ser compatible con el case.',
    ])
    doc.add_paragraph(
        'Si se detecta una incompatibilidad, el sistema muestra alertas visuales y puede bloquear '
        'la selección del componente incompatible.'
    )

    doc.add_heading('3.8 Diagrama de compatibilidad', level=2)
    doc.add_paragraph(
        'El diagrama de compatibilidad muestra visualmente las conexiones entre los componentes seleccionados '
        'y resalta posibles conflictos. Acceda mediante el botón de diagrama en la sección de resumen.'
    )

    doc.add_heading('3.9 Sección de Extras', level=2)
    doc.add_paragraph(
        'El paso "Otros" permite agregar componentes opcionales organizados en subsecciones con acordeones:'
    )
    add_styled_table(doc,
        ['Subsección', 'Categorías incluidas'],
        [
            ['Periféricos', 'Mouse, Teclado, Mousepad, Webcam'],
            ['Audio', 'Auricular, Parlante'],
            ['Software', 'Windows, Office, Antivirus'],
            ['Almacenamiento externo', 'Disco externo / USB'],
            ['Energía', 'UPS, Estabilizador'],
            ['Monitor', 'Monitores'],
            ['Refrigeración', 'Cooler aire, Cooler líquido'],
            ['Conectividad', 'Adaptadores y tarjetas de red'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Cada extra tiene controles de cantidad (+/−) para agregar múltiples unidades. '
        'El subtotal por subsección se muestra en tiempo real.'
    )

    doc.add_heading('3.10 Embalaje y Flete', level=2)
    doc.add_paragraph(
        'Estas secciones solo están disponibles para administradores autenticados:'
    )
    add_bullet_list(doc, [
        'Embalaje: active el interruptor para incluir embalaje básico ($20) o avanzado ($30). Los precios son editables.',
        'Flete: active el interruptor para incluir costo de transporte. El precio por defecto es $20 y es editable.'
    ])
    add_note(doc, 'Estos costos se suman al total de la cotización y se reflejan en el resumen financiero.')

    doc.add_heading('3.11 Comparador de productos', level=2)
    doc.add_paragraph(
        'El comparador permite evaluar hasta 3 productos lado a lado. Para usarlo:'
    )
    add_numbered_list(doc, [
        'Haga clic en "Comparar" en la tarjeta de un producto.',
        'Repita con hasta 2 productos más.',
        'El panel lateral mostrará las especificaciones y precios comparados.',
        'Use "Limpiar comparador" para reiniciar.'
    ])

    doc.add_heading('3.12 Analizador de presupuesto', level=2)
    doc.add_paragraph(
        'El analizador de presupuesto muestra una representación visual de cómo se distribuye el gasto '
        'entre los componentes seleccionados. Le permite identificar si un componente concentra '
        'demasiado del presupuesto y optimizar la configuración.'
    )

    doc.add_heading('3.13 Balance final', level=2)
    doc.add_paragraph(
        'El balance final muestra el desglose financiero completo de la configuración:'
    )
    add_bullet_list(doc, [
        'Subtotal de componentes principales.',
        'Subtotal de extras.',
        'Costo de embalaje (si está activo).',
        'Costo de flete (si está activo).',
        'Margen de ganancia aplicado.',
        'IGV.',
        'Total en USD y PEN.'
    ])

    doc.add_heading('3.14 Resumen financiero admin', level=2)
    doc.add_paragraph(
        'Visible solo para administradores. Muestra el desglose detallado con costo base, '
        'margen, IGV, y totales en ambas monedas. Permite verificar la rentabilidad de la cotización.'
    )

    doc.add_heading('3.15 Datos del cliente', level=2)
    doc.add_paragraph('Los datos requeridos para generar la cotización dependen del rol:')
    add_styled_table(doc,
        ['Rol', 'Datos requeridos'],
        [
            ['Invitado / Usuario', 'Nombre completo (obligatorio), Correo electrónico (obligatorio), Teléfono (opcional pero recomendado)'],
            ['Admin', 'Correo del cliente (opcional), Teléfono (opcional)'],
            ['Usuario registrado', 'Ninguno — el backend usa los datos de su cuenta automáticamente'],
        ]
    )

    doc.add_heading('3.16 Generar cotización', level=2)
    doc.add_paragraph(
        'Una vez completados todos los pasos obligatorios y los datos del cliente, haga clic en '
        '"Generar cotización". El sistema creará un registro con código de ticket único (formato NSG-AÑO-NÚMERO) '
        'y fecha de validez. Podrá descargar:'
    )
    add_bullet_list(doc, [
        'PDF Comercial: cotización formal para el cliente con precios, componentes y datos de la empresa.',
        'PDF Técnico (solo admin): versión detallada con especificaciones técnicas completas.',
        'Excel (solo admin): hoja de cálculo con desglose de componentes y precios.',
    ])

    doc.add_heading('3.17 Favoritos', level=2)
    doc.add_paragraph(
        'Mientras navega el cotizador, puede marcar productos como favoritos con el ícono de corazón (❤). '
        'Los favoritos se guardan en su perfil y puede consultarlos desde la página de Perfil. '
        'Para quitar un favorito, haga clic en el ícono de eliminación (✕) en la tarjeta del producto.'
    )

    doc.add_heading('3.18 Compartir configuración', level=2)
    doc.add_paragraph(
        'Puede compartir su configuración con otras personas generando un enlace único. '
        'El enlace codifica los IDs de los componentes seleccionados en base64 como parámetro de URL. '
        'Al abrir el enlace, el sistema carga automáticamente los componentes y redirige al cotizador. '
        'Si algún componente ya no está disponible, se informa con un mensaje de advertencia.'
    )

    doc.add_heading('3.19 Asistente de IA', level=2)
    doc.add_paragraph(
        'El botón flotante del Asistente IA (ícono de estrella) abre un panel de chat donde puede:'
    )
    add_bullet_list(doc, [
        'Consultar recomendaciones de configuración según presupuesto o uso.',
        'Preguntar sobre compatibilidad entre componentes.',
        'Recibir sugerencias de alternativas para componentes seleccionados.',
        'Usar respuestas rápidas predefinidas para consultas comunes.',
        'Solicitar conexión con un asesor humano vía WhatsApp.',
    ])
    doc.add_paragraph(
        'El asistente funciona en tres modos configurables por el administrador:'
    )
    add_bullet_list(doc, [
        'Pipeline Multi-Agente NVIDIA: usa tres modelos especializados (clasificador, embeddings, reranker).',
        'Uni-modelo NVIDIA: un solo modelo NVIDIA para todas las respuestas.',
        'Uni-modelo Gemini: un modelo Google Gemini para todas las respuestas.'
    ])
    add_tip(doc, 'El asistente recuerda su historial de conversación anterior y le ofrece continuarla o iniciar una nueva.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. HISTORIAL DE COTIZACIONES
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('4. Historial de Cotizaciones', level=1)

    doc.add_heading('4.1 Rol usuario', level=2)
    doc.add_paragraph(
        'Si tiene rol de usuario, el historial se carga automáticamente al ingresar a la página. '
        'Se muestran todas sus cotizaciones previas con sus estados, montos y opciones de descarga. '
        'Si no tiene cotizaciones aún, se muestra un estado vacío con enlace directo al cotizador.'
    )

    doc.add_heading('4.2 Rol administrador', level=2)
    doc.add_paragraph(
        'Los administradores pueden buscar cotizaciones por correo electrónico del cliente. '
        'Además, la página muestra una lista de clientes registrados con su número de cotizaciones. '
        'Haga clic en un cliente para cargar su historial automáticamente.'
    )

    doc.add_heading('4.3 Filtro por estado', level=2)
    add_styled_table(doc,
        ['Estado', 'Badge', 'Significado'],
        [
            ['Pendiente', 'Amarillo', 'Cotización activa, aún no reclamada por el cliente.'],
            ['Completada', 'Verde', 'El cliente reclamó la cotización y se completó la entrega.'],
            ['Reclamada', 'Azul', 'El cliente ha reclamado la cotización.'],
            ['Vencida', 'Rojo', 'La cotización superó su fecha de validez y ya no es válida.'],
        ]
    )

    doc.add_heading('4.4 Descargas disponibles', level=2)
    add_bullet_list(doc, [
        'PDF Comercial: disponible para todos los roles.',
        'PDF Técnico: solo administradores.',
        'Excel: solo administradores.',
    ])
    add_note(doc, 'Si la cotización está vencida, el sistema bloquea la generación de PDF con un mensaje de error.')

    doc.add_heading('4.5 Validar ticket', level=2)
    doc.add_paragraph(
        'Los administradores pueden hacer clic en "Validar" junto a una cotización para abrir el Validador '
        'con el código de ticket pre-cargado y revisar precios actuales vs originales.'
    )

    doc.add_heading('4.6 Notificaciones', level=2)
    doc.add_paragraph(
        'La columna "Notificación" muestra el estado del envío de correo al cliente:'
    )
    add_bullet_list(doc, [
        'Enviada (verde): correo entregado correctamente.',
        'Fallida (rojo): no se pudo enviar el correo.',
        'Pendiente (amarillo): correo en proceso de envío.',
        'Sin envíos: aún no se ha notificado al cliente.'
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. VALIDADOR DE COTIZACIONES
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('5. Validador de Cotizaciones', level=1)
    doc.add_paragraph(
        'El Validador permite verificar el estado, precios y disponibilidad de una cotización antes de confirmar la entrega. '
        'Solo está disponible para administradores autenticados.'
    )

    doc.add_heading('5.1 Buscar ticket', level=2)
    doc.add_paragraph(
        'Ingrese el código de ticket en el campo de búsqueda (formato NSG-AÑO-NÚMERO, ej: NSG-2026-0001) '
        'y haga clic en "Validar ticket". El sistema consultará la cotización y mostrará el detalle completo.'
    )

    doc.add_heading('5.2 Auto-validación desde URL', level=2)
    doc.add_paragraph(
        'Si accede al Validador con un parámetro de URL (?ticket=NSG-2026-XXXX), el sistema ejecuta '
        'la validación automáticamente sin necesidad de hacer clic. Esto es útil para enlaces directos '
        'desde el Historial.'
    )

    doc.add_heading('5.3 Detalle de validación', level=2)
    doc.add_paragraph('Al encontrar una cotización válida, se muestran 5 tarjetas de resumen:')
    add_bullet_list(doc, [
        'Ticket: código y estado actual.',
        'Precio original: monto al momento de emisión con fecha.',
        'Subtotal neto: sin IGV, con monto de IGV desglosado.',
        'Precio actual: recalculado con precios vigentes del catálogo.',
        'Diferencia total: variación entre precio original y actual (verde si bajó, amarillo si subió).'
    ])

    doc.add_heading('5.4 Diferencia de precio por componente', level=2)
    doc.add_paragraph(
        'La tabla de componentes muestra columna "Original" (precio al emitir), "Actual" (precio vigente) '
        'y "Diferencia" (variación unitaria). Si no hay cambios, muestra "Sin cambio".'
    )

    doc.add_heading('5.5 Disponibilidad actual', level=2)
    doc.add_paragraph('Cada componente muestra su estado de disponibilidad:')
    add_bullet_list(doc, [
        'En stock (verde): unidades disponibles.',
        'A pedido (amarillo): disponible con tiempo de entrega.',
        'No disponible (rojo): producto agotado o descontinuado.'
    ])

    doc.add_heading('5.6 Marcar como completada', level=2)
    doc.add_paragraph(
        'Si la cotización está en estado "Pendiente", puede hacer clic en "Marcar como completada" para '
        'actualizar su estado. Se solicita confirmación en un modal. Verifique que el cliente esté presente '
        'antes de confirmar.'
    )

    doc.add_heading('5.7 Notificar equipo listo', level=2)
    doc.add_paragraph(
        'Haga clic en "Notificar equipo listo" para enviar un correo al cliente informándole que su '
        'equipo está disponible para recojo. Disponible para cotizaciones en estado Pendiente o Completada.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. ADMINISTRACIÓN DE PRODUCTOS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('6. Administración de Productos', level=1)
    doc.add_paragraph(
        'La página de Gestión de Productos permite administrar todo el catálogo del sistema. '
        'Solo está disponible para administradores autenticados.'
    )

    doc.add_heading('6.1 Vistas de productos', level=2)
    add_bullet_list(doc, [
        'Vista compacta: tabla con columnas ID, Producto, Categoría, Precio, Stock, Pedido, Estado IA y Acciones. '
        'Incluye buscador y ordenamiento por columna.',
        'Vista detallada: tarjetas con especificaciones técnicas completas, descripción, badges de estado y precios.'
    ])

    doc.add_heading('6.2 Filtros', level=2)
    doc.add_paragraph('Puede filtrar productos por:')
    add_bullet_list(doc, [
        'Categoría: seleccionar del menú desplegable (procesador, placa_madre, ram, etc.).',
        'Estado IA: Todos, Datos CSV, Completado IA, IA Falló, Pendiente IA, Sin enriquecimiento.',
        'Búsqueda libre: escribe en el buscador para filtrar por nombre, categoría o especificación.'
    ])

    doc.add_heading('6.3 Crear producto nuevo', level=2)
    doc.add_paragraph(
        'Haga clic en "Nuevo producto" para abrir el formulario modal. Los campos varían según la categoría seleccionada:'
    )
    add_styled_table(doc,
        ['Categoría', 'Campos específicos'],
        [
            ['Procesador', 'Socket, Arquitectura, Núcleos, Hilos, Frecuencia base/boost (GHz), TDP (W), Gráficos integrados'],
            ['Placa Madre', 'Socket, Chipset, Form factor, Tipo RAM, RAM máxima (GB), Slots RAM, PCIe, M.2'],
            ['RAM', 'Tipo RAM, Capacidad (GB), Velocidad (MHz), Latencia, Módulos, Cantidad, RGB'],
            ['Almacenamiento', 'Tipo, Capacidad (GB), Interfaz, Form factor, Lectura/Escritura (MB/s), NVMe Gen'],
            ['GPU', 'Chipset, VRAM (GB), Tipo VRAM, Bus (bits), Boost (MHz), TDP (W), Longitud (mm), Fuente recomendada (W)'],
            ['Fuente', 'Potencia (W), Certificación, Modularidad, Form factor, Conectores PCIe/SATA'],
            ['Case', 'Form factor, Compatibilidad placa, GPU máxima (mm), Cooler máximo (mm), Ventiladores, Color, Panel lateral'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Campos comunes a todas las categorías:')
    add_bullet_list(doc, [
        'Nombre (obligatorio)',
        'Categoría (obligatorio)',
        'Precio base (obligatorio)',
        'Stock (obligatorio)',
        'Disponible a pedido + Tiempo de entrega en días',
        'Descripción técnica',
        'URL de imagen',
    ])

    doc.add_heading('6.4 Editar producto', level=2)
    doc.add_paragraph(
        'Haga clic en "Editar" en la fila o tarjeta del producto. El modal se pre-llenará con los datos actuales. '
        'Modifique los campos necesarios y guarde los cambios.'
    )

    doc.add_heading('6.5 Eliminar producto', level=2)
    doc.add_paragraph(
        'Haga clic en "Eliminar" y confirme en el modal de confirmación. Esta acción es irreversible.'
    )
    add_warning(doc, 'La eliminación de un producto no se puede deshacer. Verifique que el producto no esté en cotizaciones activas antes de eliminarlo.')

    doc.add_heading('6.6 Historial de precios', level=2)
    doc.add_paragraph(
        'Los productos que han tenido cambios de precio muestran un badge "H" azul y el precio anterior tachado '
        'debajo del precio actual. Haga clic en el badge para ver el detalle del historial.'
    )

    doc.add_heading('6.7 Estado de enriquecimiento IA', level=2)
    add_styled_table(doc,
        ['Badge', 'Estado', 'Significado'],
        [
            ['CSV', 'Datos CSV', 'Producto importado desde CSV, sin datos técnicos enriquecidos por IA.'],
            ['Completado IA', 'Datos completados por IA', 'La IA completó automáticamente los campos técnicos.'],
            ['IA Falló', 'Fallo de IA', 'La IA no pudo completar los datos. Puede editar manualmente o reintentar.'],
            ['Pendiente IA', 'Pendiente', 'El enriquecimiento IA está en cola o en proceso.'],
            ['Sin enriquecimiento', 'No aplica', 'Producto creado manualmente, no requiere enriquecimiento IA.'],
        ]
    )

    doc.add_heading('6.8 Banner de productos pendientes', level=2)
    doc.add_paragraph(
        'Si hay productos pendientes de enriquecimiento IA, aparece un banner amarillo en la parte superior '
        'con el conteo y un enlace "Ver estado →" que redirige a la página de Importar CSV.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. IMPORTAR CSV
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('7. Importar CSV', level=1)
    doc.add_paragraph(
        'La importación CSV permite actualizar el catálogo de productos masivamente desde un archivo .csv. '
        'Solo disponible para administradores.'
    )

    doc.add_heading('7.1 Cargar archivo', level=2)
    doc.add_paragraph('Puede cargar el archivo de dos formas:')
    add_numbered_list(doc, [
        'Arrastrar y soltar: arrastre el archivo .csv directamente sobre la zona punteada.',
        'Selección manual: haga clic en la zona para abrir el selector de archivos del sistema.',
    ])
    add_note(doc, 'Solo se permiten archivos con extensión .csv.')

    doc.add_heading('7.2 Previsualización', level=2)
    doc.add_paragraph(
        'Al cargar el archivo, el sistema muestra una previsualización con las 10 primeras filas '
        'en una tabla con scroll horizontal y vertical. Esto le permite verificar que las columnas '
        'y datos estén correctos antes de importar.'
    )

    doc.add_heading('7.3 Ejecutar importación', level=2)
    doc.add_paragraph(
        'Haga clic en "Importar" para iniciar el proceso. El botón muestra un indicador de carga '
        'mientras se procesa. Puede cancelar con el botón "Cancelar".'
    )

    doc.add_heading('7.4 Resultado de la importación', level=2)
    doc.add_paragraph('Al completar, el sistema muestra 5 contadores:')
    add_styled_table(doc,
        ['Contador', 'Color', 'Significado'],
        [
            ['Insertados', 'Verde', 'Productos nuevos agregados al catálogo.'],
            ['Actualizados', 'Azul', 'Productos existentes con datos actualizados.'],
            ['Omitidos', 'Amarillo', 'Productos que no cumplieron criterios de importación.'],
            ['Errores', 'Rojo', 'Filas con datos inválidos que no se procesaron.'],
            ['Pendientes IA', 'Morado', 'Productos encolados para enriquecimiento automático por IA.'],
        ]
    )

    doc.add_heading('7.5 Detalle de errores', level=2)
    doc.add_paragraph(
        'Si hubo errores, se muestran debajo de los contadores con el número de fila y el mensaje de error específico. '
        'Revise estos errores para corregir su archivo CSV.'
    )

    doc.add_heading('7.6 Estado del enriquecimiento IA', level=2)
    doc.add_paragraph(
        'Después de la importación, aparece una sección "Estado del enriquecimiento IA" que muestra:'
    )
    add_bullet_list(doc, [
        'Indicador animado "Enriqueciendo datos con IA..." mientras el proceso está activo.',
        'Contadores: Pendientes, Completados, Fallidos.',
        'Última actualización (fecha y hora).',
    ])
    doc.add_paragraph(
        'Esta sección se actualiza automáticamente cada 10 segundos mientras haya productos pendientes.'
    )

    doc.add_heading('7.7 Reintentar productos fallidos', level=2)
    doc.add_paragraph(
        'Si hay productos con estado "Fallido", aparece un botón "Reintentar productos fallidos" '
        'que vuelve a encolar esos productos para un nuevo intento de enriquecimiento IA.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. CONFIGURACIÓN DEL SISTEMA
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('8. Configuración del Sistema', level=1)
    doc.add_paragraph(
        'La página de Configuración permite gestionar parámetros globales del sistema y monitorear el uso de IA. '
        'Solo disponible para administradores.'
    )

    doc.add_heading('8.1 Parámetros financieros', level=2)
    doc.add_paragraph('Configure los valores que afectan el cálculo de precios de todas las cotizaciones:')
    add_bullet_list(doc, [
        'Margen de ganancia (%): porcentaje que se suma al costo base. Rango: 0-100.',
        'Tasa de IGV (%): impuesto general a las ventas. Rango: 0-100.',
        'Tipo de cambio USD/PEN: valor de conversión de dólares a soles.',
    ])
    doc.add_paragraph(
        'Se muestran cálculos de ejemplo para tres precios base (USD 1000, 2500, 5000) con su referencia en PEN.'
    )

    doc.add_heading('8.2 Modo de tipo de cambio', level=2)
    doc.add_paragraph('El sistema ofrece dos modos para el tipo de cambio:')
    add_bullet_list(doc, [
        'Manual: usted ingresa el valor del tipo de cambio directamente. El campo es editable.',
        'Automático (API): el sistema consulta automáticamente el tipo de cambio actual desde una API externa. '
        'El valor se muestra en solo lectura con la fecha de última actualización.',
    ])
    add_note(doc, 'Cambie el modo con el selector segmentado y guarde con el botón "Guardar modo".')

    doc.add_heading('8.3 Forzar actualización', level=2)
    doc.add_paragraph(
        'En modo automático, use el botón "Actualizar tipo de cambio" para forzar una nueva consulta a la API. '
        'Si la API falla, se muestra una advertencia amarilla indicando que se está usando un valor de respaldo.'
    )

    doc.add_heading('8.4 Configuración activa', level=2)
    doc.add_paragraph(
        'Un recuadro muestra siempre la configuración actualmente en vigor: Margen, IGV y Tipo de cambio. '
        'Esto le permite comparar antes de guardar cambios nuevos.'
    )

    doc.add_heading('8.5 Métricas de IA', level=2)
    doc.add_paragraph('Monitoree el uso del asistente de IA:')
    add_bullet_list(doc, [
        'Llamadas: total de consultas acumuladas al asistente.',
        'Costo estimado: gasto aproximado en USD por uso de APIs externas.',
        'Promedio tokens: tokens promedio consumidos por consulta.',
    ])

    doc.add_heading('8.6 Configuración del asistente IA', level=2)
    doc.add_paragraph('Seleccione el modo de operación del asistente:')
    add_styled_table(doc,
        ['Modo', 'Descripción', 'Campos requeridos'],
        [
            ['Pipeline Multi-Agente NVIDIA', 'Tres modelos especializados: clasificador, embeddings y reranker.', 'Modelo clasificador, Modelo embeddings, Modelo reranker'],
            ['Uni-modelo NVIDIA', 'Un único modelo NVIDIA para todas las respuestas.', 'Modelo NVIDIA'],
            ['Uni-modelo Gemini', 'Un modelo Google Gemini para todas las respuestas.', 'Modelo Gemini'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Los campos requeridos cambian dinámicamente según el modo seleccionado. '
        'Guarde con el botón "Guardar configuración de IA".'
    )

    doc.add_heading('8.7 Claves API de IA', level=2)
    doc.add_paragraph(
        'Configure las claves API para los servicios de IA:'
    )
    add_bullet_list(doc, [
        'Clave API de Gemini: formato AIza... Indicador "Configurada" o "No configurada".',
        'Clave API de NVIDIA: formato nvapi-... Indicador "Configurada" o "No configurada".',
    ])
    doc.add_paragraph(
        'Las claves se almacenan encriptadas en la base de datos. Puede mostrar/ocultar el valor con el ícono de ojo.'
    )

    doc.add_heading('8.8 Mantenimiento', level=2)
    doc.add_paragraph(
        'La sección "Mantenimiento (Testeo)" contiene acciones destructivas para desarrollo:'
    )
    add_warning(doc, '"Limpiar todo el catálogo" elimina TODOS los productos de las 23 tablas del sistema. Esta acción es irreversible y solo debe usarse en ambientes de prueba.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. DASHBOARD
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('9. Dashboard', level=1)
    doc.add_paragraph(
        'El Dashboard muestra métricas operativas del negocio en tiempo real. Solo disponible para administradores.'
    )

    doc.add_heading('9.1 Métricas principales', level=2)
    add_bullet_list(doc, [
        'Cotizaciones hoy: total de cotizaciones emitidas en el día actual (pendientes y completadas).',
        'Cotizaciones semana: total de cotizaciones en los últimos 7 días.',
        'Ingresos estimados hoy: monto total estimado de cotizaciones activas del día.',
        'Ingresos estimados semana: monto total estimado de los últimos 7 días.',
    ])

    doc.add_heading('9.2 Productos más cotizados', level=2)
    doc.add_paragraph(
        'Un gráfico de barras muestra el Top 5 de productos más cotizados en los últimos 7 días. '
        'Debajo del gráfico hay una tabla accesible con nombre, categoría y número de apariciones.'
    )

    doc.add_heading('9.3 Actualización', level=2)
    doc.add_paragraph(
        'Use el botón "Actualizar" para recargar las métricas manualmente.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. PERFIL Y FAVORITOS
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('10. Perfil y Favoritos', level=1)

    doc.add_heading('10.1 Información del perfil', level=2)
    doc.add_paragraph(
        'La página de Perfil muestra su nombre de usuario y correo. Es el punto de acceso para gestionar sus favoritos.'
    )

    doc.add_heading('10.2 Lista de favoritos', level=2)
    doc.add_paragraph(
        'Se muestran todos los productos marcados como favoritos con:'
    )
    add_bullet_list(doc, [
        'Nombre del producto.',
        'Categoría.',
        'Descripción técnica (truncada).',
        'Especificaciones clave (socket, tipo RAM, potencia).',
        'Precio en la moneda seleccionada y referencia en la otra moneda.',
        'Fecha en que fue agregado a favoritos.',
    ])

    doc.add_heading('10.3 Quitar favorito', level=2)
    doc.add_paragraph(
        'Haga clic en el botón ✕ de la tarjeta para quitar el producto de favoritos. '
        'La eliminación es inmediata (actualización optimista). Si hay error, el producto se restaura automáticamente.'
    )

    doc.add_heading('10.4 Ir al cotizador', level=2)
    doc.add_paragraph(
        'Si tiene favoritos, aparece un banner inferior "¿Listo para armar tu configuración?" '
        'con un botón para ir directamente al cotizador.'
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. NAVEGACIÓN Y ACCESIBILIDAD
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('11. Navegación y Accesibilidad', level=1)

    doc.add_heading('11.1 Barra de navegación', level=2)
    doc.add_paragraph('La barra lateral (escritorio) y los tabs inferiores (móvil) permiten acceder a:')
    add_styled_table(doc,
        ['Sección', 'Visible para', 'Icono'],
        [
            ['Cotizador', 'Todos', 'Estrella'],
            ['Historial', 'Autenticados', 'Reloj'],
            ['Iniciar sesión', 'Invitados', 'Candado'],
            ['Crear cuenta', 'Invitados', 'Persona+'],
            ['Validador', 'Admin', 'Lupa'],
            ['Productos (Admin)', 'Admin', 'Caja'],
            ['Configuración', 'Admin (solo lateral)', 'Sliders'],
            ['Importar CSV', 'Admin (solo lateral)', 'Subida'],
        ]
    )

    doc.add_heading('11.2 Modo oscuro', level=2)
    doc.add_paragraph(
        'El sistema soporta modo claro y oscuro. Puede cambiar entre modos desde el selector de tema '
        'en la barra de navegación. La preferencia se guarda en el navegador.'
    )

    doc.add_heading('11.3 Selector de moneda', level=2)
    doc.add_paragraph(
        'Use el selector de moneda (USD/PEN) para ver todos los precios del sistema en la moneda preferida. '
        'La moneda alternativa se muestra como referencia entre paréntesis. '
        'La conversión usa el tipo de cambio configurado en el sistema.'
    )

    doc.add_heading('11.4 Notificaciones toast', level=2)
    doc.add_paragraph(
        'El sistema muestra notificaciones temporales (toast) en la esquina superior para confirmar acciones '
        '(éxito, error, advertencia, información). Las notificaciones desaparecen automáticamente.'
    )

    doc.add_heading('11.5 Panel de notificaciones', level=2)
    doc.add_paragraph(
        'El ícono de campana muestra notificaciones del sistema con badge de conteo. '
        'Haga clic para abrir el panel lateral con las notificaciones recientes.'
    )

    doc.add_heading('11.6 Accesibilidad', level=2)
    doc.add_paragraph(
        'El sistema cumple con los estándares WCAG AA:'
    )
    add_bullet_list(doc, [
        'Contraste de colores adecuado en modo claro y oscuro.',
        'Navegación completa por teclado con indicadores de foco visibles.',
        'Targets táctiles mínimos de 44px para uso en dispositivos móviles.',
        'Etiquetas ARIA correctas en todos los elementos interactivos.',
        'Animaciones que respetan la preferencia "movimiento reducido" del sistema operativo.',
    ])

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. REQUISITOS E INSTALACIÓN
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('12. Requisitos e Instalación', level=1)

    doc.add_heading('12.1 Requisitos del sistema', level=2)
    add_styled_table(doc,
        ['Componente', 'Requisito mínimo'],
        [
            ['Node.js', 'v18 o superior'],
            ['PostgreSQL', 'v14 o superior'],
            ['Navegador', 'Chrome, Firefox, Safari o Edge (últimas 2 versiones)'],
            ['RAM', '2 GB mínimo (4 GB recomendado)'],
            ['Disco', '500 MB para la aplicación + espacio para base de datos'],
        ]
    )

    doc.add_heading('12.2 Variables de entorno', level=2)
    doc.add_paragraph('Backend (.env):')
    add_styled_table(doc,
        ['Variable', 'Descripción', 'Ejemplo'],
        [
            ['DB_HOST', 'Host de PostgreSQL', 'localhost'],
            ['DB_PORT', 'Puerto de PostgreSQL', '5432'],
            ['DB_NAME', 'Nombre de la base de datos', 'nsg_cotizaciones'],
            ['DB_USER', 'Usuario de PostgreSQL', 'postgres'],
            ['DB_PASSWORD', 'Contraseña de PostgreSQL', '(su contraseña)'],
            ['JWT_SECRET', 'Clave secreta para tokens JWT', 'mínimo 32 caracteres aleatorios'],
            ['ENCRYPTION_KEY', 'Clave para encriptar API keys', '64 caracteres hexadecimales'],
            ['PORT', 'Puerto del servidor backend', '3000'],
            ['FRONTEND_URL', 'URL del frontend', 'http://localhost:5173'],
            ['SMTP_HOST', 'Servidor SMTP para correos', 'smtp.gmail.com'],
            ['SMTP_USER', 'Correo remitente', 'tu_correo@gmail.com'],
            ['SMTP_PASS', 'Clave de aplicación', '(clave de app de Google)'],
            ['APIS_NET_TOKEN', 'Token API tipo de cambio', 'sk_...'],
            ['GEMINI_API_KEY', 'Clave API Google Gemini', 'AIza...'],
            ['NVIDIA_API_KEY', 'Clave API NVIDIA', 'nvapi-...'],
            ['WHATSAPP_NUMERO_ASESOR', 'WhatsApp del asesor', '51999999999'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('Frontend (.env):')
    add_styled_table(doc,
        ['Variable', 'Descripción', 'Ejemplo'],
        [
            ['VITE_API_URL', 'URL base del API backend', 'http://localhost:3000/api'],
            ['VITE_API_TIMEOUT_MS', 'Timeout de peticiones (ms)', '300000'],
        ]
    )

    doc.add_heading('12.3 Comandos de instalación y ejecución', level=2)
    doc.add_paragraph('Configuración inicial:')
    add_numbered_list(doc, [
        'Clonar el repositorio e ingresar al directorio del proyecto.',
        'Backend: cd backend → copiar .env.example como .env → completar variables → npm install → npm run dev.',
        'Frontend: cd frontend → copiar .env.example como .env → npm install → npm run dev.',
        'Acceder a http://localhost:5173 en el navegador.',
    ])
    doc.add_paragraph('Comandos útiles:')
    add_styled_table(doc,
        ['Comando', 'Ubicación', 'Descripción'],
        [
            ['npm run dev', 'backend/', 'Inicia servidor en modo desarrollo con recarga automática.'],
            ['npm start', 'backend/', 'Inicia servidor en modo producción.'],
            ['npm test', 'backend/', 'Ejecuta pruebas unitarias.'],
            ['npm run seed', 'backend/', 'Carga datos de prueba en la base de datos.'],
            ['npm run dev', 'frontend/', 'Inicia Vite en modo desarrollo.'],
            ['npm run build', 'frontend/', 'Genera build de producción.'],
        ]
    )

    doc.add_heading('12.4 Roles del sistema', level=2)
    add_styled_table(doc,
        ['Rol', 'Acceso al cotizador', 'Admin productos', 'Importar CSV', 'Configuración', 'Validador', 'Dashboard'],
        [
            ['Invitado', 'Sí (sin precios)', 'No', 'No', 'No', 'No', 'No'],
            ['Usuario', 'Sí', 'No', 'No', 'No', 'No', 'No'],
            ['Admin', 'Sí', 'Sí', 'Sí', 'Sí', 'Sí', 'Sí'],
        ]
    )
    add_note(doc, 'Los invitados ven "Inicia sesión para ver precio" en lugar de los precios reales.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 13. GLOSARIO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('13. Glosario', level=1)
    add_styled_table(doc,
        ['Término', 'Definición'],
        [
            ['Cotización', 'Documento formal con precio, componentes y validez temporal para un equipo de cómputo.'],
            ['Ticket', 'Código único asignado a cada cotización (formato NSG-AÑO-NÚMERO).'],
            ['Margen de ganancia', 'Porcentaje que se suma al costo base del producto para determinar el precio de venta.'],
            ['IGV', 'Impuesto General a las Ventas (18% en Perú por defecto).'],
            ['Tipo de cambio', 'Valor de conversión entre dólares americanos (USD) y soles peruanos (PEN).'],
            ['Enriquecimiento IA', 'Proceso automático donde la inteligencia artificial completa datos técnicos de productos importados por CSV.'],
            ['Socket', 'Conector de la placa madre donde se instala el procesador. Debe coincidir entre ambos.'],
            ['Form factor', 'Tamaño estándar de la placa madre (ATX, Micro-ATX, Mini-ITX). Determina compatibilidad con el case.'],
            ['TDP', 'Thermal Design Power — potencia térmica máxima que disipa un componente.'],
            ['VRAM', 'Memoria de video dedicada de la tarjeta gráfica.'],
            ['NVMe', 'Non-Volatile Memory Express — protocolo de alta velocidad para almacenamiento M.2.'],
            ['iGPU', 'Gráficos integrados en el procesador (Intel UHD, AMD Radeon Graphics).'],
            ['Pipeline Multi-Agente', 'Arquitectura de IA que encadena tres modelos especializados para generar respuestas más precisas.'],
            ['QR Code', 'Código QR incluido en los PDF de cotización para validación rápida en tienda.'],
            ['Toast', 'Notificación temporal no intrusiva que aparece brevemente en la interfaz.'],
            ['WCAG AA', 'Web Content Accessibility Guidelines nivel AA — estándar de accesibilidad web.'],
        ]
    )

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 14. SOPORTE Y CONTACTO
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_heading('14. Soporte y Contacto', level=1)
    doc.add_paragraph(
        'Para soporte técnico, consultas o sugerencias sobre el sistema NSG Cotizador, contacte al equipo de desarrollo:'
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('NSG Latinoamérica E.I.R.L.')
    run.bold = True
    doc.add_paragraph('WhatsApp del asesor: configurable en el sistema (variable WHATSAPP_NUMERO_ASESOR).')
    doc.add_paragraph()
    doc.add_paragraph(
        'Para reportar errores del sistema, incluya en su mensaje:'
    )
    add_bullet_list(doc, [
        'Descripción del problema.',
        'Pasos para reproducirlo.',
        'Código de ticket (si aplica).',
        'Captura de pantalla del error.'
    ])

    # ── GUARDAR ─────────────────────────────────────────────────────────────
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Manual_Usuario_NSG_Cotizador.docx')
    doc.save(output_path)
    print(f'Manual generado exitosamente: {output_path}')
    return output_path


if __name__ == '__main__':
    generar_manual()
